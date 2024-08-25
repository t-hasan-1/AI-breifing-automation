#LIBRARIES
import json
import requests
from datetime import datetime
import re
from bs4 import BeautifulSoup
from readability import Document
from docx import Document as WordDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import urllib.parse
from datetime import datetime
from datetime import timedelta

#CONSTANTS 
API_KEY = "76a0fc9b20b74420b155777516ecda5d"
BASE_URL = "https://newsapi.org/v2/everything?"
TERMS_NATIONAL="(AI OR CHATGPT) OR (GOVERNMENT AND AI) AND (NOT INDIA NOT CHINA NOT RUSSIA NOT EUROPE)"
TERMS_WORLD="(AI OR CHATGPT) OR (GOVERNMENT AND AI) AND (CHINA OR RUSSIA OR INDIA OR EUROPE OR EU)"
TO_DATE = datetime.today().strftime('%Y-%m-%d')
FROM_DATE = (datetime.today() - timedelta(hours=14)).strftime('%Y-%m-%d')
LANGUAGE = "en"
CACHED_RESPONSE_NATIONAL = "data_national.json"
CACHED_RESPONSE_WORLD = "data_world.json"
FETCH_REMOTE = True #Pull from cached json file, or go to newsapi

#FUNCTIONS 
def main():
    if FETCH_REMOTE:
        encoded_query_national = urllib.parse.quote("".join(TERMS_NATIONAL))
        encoded_query_world = urllib.parse.quote("".join(TERMS_WORLD))
        final_url_national =  f"{BASE_URL}q={encoded_query_national}&from={FROM_DATE}&to={TO_DATE}&language={LANGUAGE}&pageSize=30&apiKey={API_KEY}"
        final_url_world =  f"{BASE_URL}q={encoded_query_world}&from={FROM_DATE}&to={TO_DATE}&language={LANGUAGE}&pageSize=30&apiKey={API_KEY}"

        response_json_national = fetch_articles(final_url_national)
        response_json_world = fetch_articles(final_url_world)
        save_to_file(response_json_national, CACHED_RESPONSE_NATIONAL)
        save_to_file(response_json_world, CACHED_RESPONSE_WORLD)
    else:
        response_json_national = load_from_file(CACHED_RESPONSE_NATIONAL)
        response_json_world = load_from_file(CACHED_RESPONSE_WORLD)

    template = load_template()

    articles_national = response_json_national.get('articles', [])
    articles_world = response_json_world.get('articles', [])

    articles_processed = 0
    while(articles_processed) < 5:
        article = articles_national[0]
        articles_national.remove(article)

        # processed is true if this article was written to output. Else try another article
        processed = format_output(article, template, articles_processed + 1)

        if processed:
            articles_processed += 1

    while(articles_processed) < 7:
        article = articles_world[0]
        articles_world.remove(article)
        processed = format_output(article, template, articles_processed + 1)

        if processed:
            articles_processed += 1

    template.save(f'FINAL_OUTPUT_{TO_DATE}.docx')

def fetch_articles(url):
    # print("FETCHING TO: " + url)
    response = requests.get(url)
    return response.json()

def save_to_output_file(data, filename):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(data)

def save_to_file(data, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def load_from_file(filename):
    with open(filename, 'r', encoding='utf-8') as file:
        return json.load(file)

def extract_first_7_sentences(text):
    # Use a more robust regex to split sentences
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|!)\s', text)
    return ' '.join(sentences[:7])

def check_manual_review_article(url, articleIndex):
    print(f"Please review the following URL for article #{articleIndex}: {url}")
    
    # Ask for user confirmation
    response = input("Is the URL acceptable? (yes/no): ").strip().lower()
    if response == 'yes':
        return True
    elif response == 'no':
        return False

def format_output(article, template, article_index):
    source = article.get('source', {}).get('name', 'Unknown Source') or "Unknown Source"
    title = article.get('title', 'No Title') or "No Title"
    author = article.get('author', 'No Author') or "No Author"
    date = datetime.strptime(article.get('publishedAt', '1970-01-01T00:00:00Z'), "%Y-%m-%dT%H:%M:%SZ").strftime('%m/%d/%Y')
    description = article.get('description', 'No Description')
    url = article.get('url', '#')


    # Scrape the full article content
    content = scrape_article_content(url)
    snippet = extract_first_7_sentences(content)

    # If this is a bad article, try another
    is_bad_article = source == "Unknown Source" or title == "No Title" or author == "No Author" or url == "#" or content == "No content available" 
    if is_bad_article:
        return False
    if not check_manual_review_article(url, article_index):
        return False

    # Modify the tempalte by replacing keywords
    replace_template_keyword(template, f"{article_index}_SOURCE", source)
    replace_template_keyword(template, f"{article_index}_TITLE", title)
    replace_template_keyword(template, f"{article_index}_AUTHOR", author)
    replace_template_keyword(template, f"{article_index}_CONTENT", snippet)
    replace_template_keyword(template, f"{article_index}_DATE", date)
    replace_template_keyword(template, f"{article_index}_HYPERLINK", url)
    return True

def scrape_article_content(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        html = response.text
        doc = Document(html)
        soup = BeautifulSoup(doc.summary(), 'html.parser')
        return soup.get_text()
    except Exception as e:
        return "No content available"

def replace_template_keyword(doc, keyword, replacement_value):
    replaced = False
    for paragraph in doc.paragraphs:
        if keyword in paragraph.text:
            for run in paragraph.runs:
                if keyword in run.text:
                    run.text = run.text.replace(keyword, replacement_value)
                    replaced = True
    
    if not replaced:
        for paragraph in doc.paragraphs:
            if keyword in paragraph.text:
                paragraph.text = paragraph.text.replace(keyword, replacement_value)

def load_template():
    TEMPLATE_PATH = "email_template.docx"
    doc = WordDocument(TEMPLATE_PATH)
    return doc      

if __name__ == "__main__":
    main()